"""
AI项目四：AI自动化数据报告系统
1.基于 Python+Pandas + 大模型 API 开发企业级自动化全流程数据报告系统
2.实现 Excel 数据读取、清洗、统计、AI 分析、Word 报告生成、自动邮件发送全流程
3.支持无人值守运行，大幅降低人工报表工作量（效率提升 90%）
4.具备日志记录、异常处理能力
5.技术栈：Python、Pandas、大模型 API、自动化办公、邮件服务、工程化部署
2026-03-22 优化内容：
1.采用 PM2 守护进程 + 定时任务 实现 7×24 小时无人值守运行
2.支持崩溃自动重启、开机自启、日志持久化、异常告警
3.真实落地场景：每日销售 / 运营 / 财务自动化报表
"""
import os
import pandas as pd
from docx import Document
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from core.logger import logger
from core.llm_client import MODEL_CONFIG,MODEL_TYPE,call_llm

# 目录创建
os.makedirs("projects/auto_report/data",exist_ok=True)
os.makedirs("projects/auto_report/output",exist_ok=True)
   
# 1.读取并清洗数据
def load_data(filepath = "projects/auto_report/data/sales.xlsx"):
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"未找到数据文件：{filepath}")
    df = pd.read_excel(filepath,engine="openpyxl")
    df = df.dropna()
    logger.info(f"数据加载完成，共{len(df)}行")
    return df
# 2.数据统计
def analyze_data(df):
    required_cols = ["销售额","城市"]
    for col in required_cols:
        if col not in df.columns:
            raise ValueError(f"数据表中缺少必要的列：{col}")
    stats = {
        "总销售额": round(df["销售额"].sum(),2),
        "平均销售额": round(df["销售额"].mean(),2),
        "最高订单": round(df["销售额"].max(),2),
        "订单总数": len(df),
        "城市数": df["城市"].nunique()
    }
    logger.info(f"统计完成：{stats}")
    return stats
# 生成Word报告
def generate_word_report(stats,analysis,output_path):
    doc = Document()
    title = doc.add_heading("销售数据自动化分析报告",0)
    title.alignment = 1

    doc.add_heading("一、数据概览",level=1)
    for k,v in stats.items():
        doc.add_paragraph(f"{k}:{v}")

    doc.add_heading("二、AI智能分析",level=1)
    doc.add_paragraph(analysis)

    doc.add_paragraph(f"\n报告生成时间：{datetime.now().strftime('%Y-%m-%d %H-%M')}")
    doc.save(output_path)
    logger.info(f"报告已生成：{output_path}")

# 发送邮件
def send_email_with_report(report_path):
    from_email = MODEL_CONFIG["email"]["email_user"]
    to_email = MODEL_CONFIG["email"]["to_email"] if MODEL_CONFIG["email"]["to_email"] else from_email
    password = MODEL_CONFIG["email"]["email_pass"]
    smtp_server = MODEL_CONFIG["email"]["email_smtp"]
    port = int(MODEL_CONFIG["email"]["email_port"] if MODEL_CONFIG["email"]["email_port"] else 465)
    if not all([from_email,password,smtp_server]):
        logger.error("邮件配置缺失，请检查环境变量")
        return
    msg = MIMEMultipart()
    msg["From"] = from_email
    msg["To"] = to_email
    msg["Subject"] = "AI自动生成 - 销售数据分析报告"
    msg.attach(MIMEText("领导您好，已自动生成今日销售数据报告，请查收。","plain","utf-8"))
    with open(report_path,"rb") as f:
        part = MIMEApplication(f.read(),Name=os.path.basename(report_path))
        part["Content-Disposition"] = f'attachment;filename="{os.path.basename(report_path)}"'
        msg.attach(part)
    try:
        server = smtplib.SMTP_SSL(smtp_server,port)
        server.login(from_email,password)
        server.sendmail(from_email,to_email,msg.as_string())
        server.quit()
        logger.info("邮件发送成功")
    except Exception as e:
        logger.error(f"邮件发送失败：{e}")

# 主流程
def main():
    logger.info("===开始执行自动化报告任务===")
    try:
        # print("请选择模型")
        # for i,m in enumerate(MODEL_TYPE):
        #     print(f"{i+1}.{m}")
        # while True:
        #     try:
        #         choice = int(input("请输入有效序号,输入其他内容均无效：\n").strip())
        #         if choice not in [1,len(MODEL_TYPE)]:
        #             print("序号不存在，请重新输入")
        #             continue
        #         model_name = MODEL_TYPE[choice - 1]
        #         logger.info(f"已选择模型：{model_name}")
        #         break
        #     except Exception:
        #         logger.error("输入无效内容，请重新输入")
        #         continue
        """
        对该项目进行优化，添加定时推送功能，对模型选择采用固定式
        """

        df = load_data()
        stats = analyze_data(df)
        prompt = f"""
        你是专业数据分析师，根据以下销售数据写一段专业、简洁、可直接用于正式报告的分析。
        数据：{stats}
        要求：
        1.分析趋势、亮点、风险
        2.给出业务建议
        3.语言正式，不使用markdown
        """
        analysis = call_llm(prompt,model_name="doubao",temperature=0.3)
        output_path = f"projects/auto_report/output/报告_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
        generate_word_report(stats,analysis,output_path)
        send_email_with_report(output_path)
        logger.info("===自动化报告任务全部完成===")
    except Exception as e:
        logger.error(f"任务执行过程中发生严重错误：{e}")

if __name__ == "__main__":
    main()

