"""
AI项目三：AI自动化办公机器人
1.基于 Python + 大模型 API 开发企业级自动化办公工具，支持 AI 生成 Word、分析 Excel、发送个性化邮件、批量文档处理
2.整合 python-docx/openpyxl/smtplib 等库，实现办公全流程自动化
3.解决企业重复办公任务痛点，提升办公效率 80% 以上
4.技术栈：Python、大模型 API、办公自动化库、批量处理、异常处理
"""
import os
import json
from docx import Document
from openpyxl import Workbook,load_workbook
import smtplib
from email.mime.text import MIMEText
from email.header import Header
from core.llm_client import call_llm,MODEL_CONFIG,MODEL_TYPE
from core.logger import logger
   
# 功能一：AI生成Word文档
def ai_generate_word(content_topic,model_name,save_path="ai_generate.docx"):
    print(f"正在生成文档：{content_topic}")
    prompt = f"""
    你是专业的办公文档写手，根据主题生成完整的文档内容。
    主题{content_topic}
    要求：
    1.包含标题、副标题、正文（至少3级）
    2.格式规范，语言专业
    3.仅输出文档内容，不要多余解释
    """
    doc_content = call_llm(prompt,model_name,temperature=0.7)
    if not doc_content or doc_content.startswith(("LLM调用失败","返回格式异常")):
        error_msg = f"文档生成失败：{doc_content}" if doc_content else "LLM调用失败"
        print(error_msg)
        logger.error(error_msg)
        return None
    # 创建Word文档
    try:
        doc = Document()
        # 添加标题
        doc.add_heading(content_topic,level=1)
        # 按段落分割内容并添加
        if "\n\n" in doc_content:
            paragraphs = doc_content.split("\n\n")
        else:
            paragraphs = doc_content.split("\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        # 保存文档
        doc.save(save_path)
        print(f"Word文档已保存：{save_path}")
        logger.info(f"Word文档已生成并保存：{save_path}")
        return save_path
    except Exception as e:
        logger.error(f"Word文档已生成并保存：{save_path}")
        return None
  
# 功能二：AI分析Excel数据
def ai_analyze_excel(model_name,file_path,save_path="ai_analysis.xlsx"):
    print(f"正在分析Excel文件：{file_path}")
    # 1.读取Excel数据
    if not os.path.exists(file_path):
        print(f"文件不存在：{file_path}")
        logger.error(f"Excel文件不存在：{file_path}")
        return None
    try:
        # 从现有路径加载xlsx文件
        wb = load_workbook(file_path)
        # 获取当前活动seet页
        ws = wb.active
        # 获取总行数
        total_row = ws.max_row
        # 动态决定读取行数，最多50条
        limit = min(total_row - 1,50)
        # 提取数据文本
        data_text = "Excle数据：\n"
        # 获取表头
        header = [cell.value for cell in ws[1]]
        data_text += f"表头：{header}\n"
        data_text += f"总行数：{total_row - 1}\n\n"
        data_text += f"部分数据样本\n"
        # 获取数据(设定限制，避免过多)
        end_row = min(limit + 1,total_row)
        for row in ws.iter_rows(min_row=2,max_row=end_row,values_only=True):
           # 过滤全空的行
           if any(cell is not None for cell in row):
               data_text += f"{row}\n"
        if total_row > limit + 1:
            data_text += "...(数据过多，仅展示部分)..."
        # AI分析数据
        prompt = f"""
        你是专业的数据分析师，分析Excel数据并生成分析报告。
        数据：{data_text}
        要求：
        1.包含数据概况、关键指标、结论建议
        2.语言简洁专业，适合放入Excel报告
        3.仅输出分析内容，不要多余解释
        """
        analysis = call_llm(prompt,model_name,temperature=0.4)
        if not analysis or analysis.startswith(("LLM调用失败", "返回格式异常")):
            error_msg = f"Excel分析失败: {analysis}" if analysis else "LLM调用失败"
            print(error_msg)
            logger.error(error_msg)
            return None
        # 生成分析报告Excel
        new_wb = Workbook()
        new_ws = new_wb.active
        new_ws.title = "数据分析报告"
        # 添加分析内容
        new_ws.cell(row=1,column=1,value="数据分析师报告")
        new_ws.cell(row=2,column=1,value="-" * 50)
        # 按行添加内容
        lines = analysis.split("\n")
        for i,line in enumerate(lines,start=3):
            if line.strip():
                new_ws.cell(row=i,column=1,value=line.strip())
        # 保存分析报告
        new_wb.save(save_path)
        print(f"数据分析报告已保存：{new_wb}")
        logger.info(f"Excel分析报告已生成并保存：{save_path}")
        return save_path
    except Exception as e:
        logger.error(f"Excel分析失败：{str(e)}")
        print(f"Excel分析失败：{str(e)}")
        return None

# 功能三：AI发送个性化邮件
def ai_send_email(model_name,to_email,subject):
    print(f"正在发送邮件：{to_email}")
    # 1.AI生成邮件内容
    prompt = f"""
    你是专业的邮件写手，根据主题生成个性化邮件内容。
    收件人：{to_email}
    主题：{subject}
    要求：
    1.包含称呼、正文、结尾、署名
    2.语言礼貌专业，格式规范
    3.仅输出邮件内容，不要多余解释
    """
    email_content = call_llm(prompt,model_name,temperature=0.6)
    if not email_content or email_content.startswith(("LLM调用失败", "返回格式异常")):
        error_msg = f"邮件内容生成失败: {email_content}" if email_content else "LLM调用失败"
        print(error_msg)
        logger.error(error_msg)
        return False
    # 2.发送邮件
    try:
        #配置邮件服务器
        smtp_obj = smtplib.SMTP_SSL(
            host = MODEL_CONFIG["email"]["email_smtp"],
            port = int(MODEL_CONFIG["email"]["email_port"]) if MODEL_CONFIG["email"]["email_port"] else 465
        )
        smtp_obj.login(
            user = MODEL_CONFIG["email"]["email_user"],
            password = MODEL_CONFIG["email"]["email_pass"]
        )
        # 构造函数
        msg = MIMEText(email_content,"plain","utf-8")
        msg["From"] = MODEL_CONFIG["email"]["email_user"]
        msg["To"] = to_email
        msg["Subject"] = Header(subject,"utf-8")
        # 发送邮件
        smtp_obj.sendmail(
            MODEL_CONFIG["email"]["email_user"],
            to_email,
            msg.as_string()
        )
        # 关闭邮件
        smtp_obj.quit()
        print(f"邮件已发送至：{to_email}")
        logger.info(f"邮件已成功发送至：{to_email}")
        return True
    except Exception as e:
        logger.error(f"邮件发送失败：{str(e)}")
        print(f"邮件发送失败：{str(e)}")
        print("提示：请检查.env中的邮箱配置是否正确")
        return False

# 功能四：批量处理文档
def batch_process_documents(model_name,topics_list,save_dir="batch_output"):
    print(f"开始批量处理，共{len(topics_list)}个主题")
    # 创建保存目录
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    # 批量生成
    results = []
    for i,topic in enumerate(topics_list,1):
        save_path = os.path.join(save_dir,f"文档{i}_{topic.replace(' ','_').replace('/','_')}.docx")
        path = ai_generate_word(topic,model_name,save_path)
        results.append({
            "topic": topic,
            "file_path": path,
            "status": "成功" if os.path.exists(path) else "失败"
        })
    # 生成处理报告
    report_path = os.path.join(save_dir,"批量处理报告.json")
    with open(report_path,"w",encoding="utf-8") as f:
        json.dump(results,f,ensure_ascii=False,indent=2)
    print(f"批量处理完成！报告已保存：{report_path}")
    logger.info(f"批量处理完成，报告保存至：{report_path}")
    return results

# 主函数：交互界面
def main():
    os.system("cls" if os.name == "nt" else "clear")
    print("=" * 60)
    print("AI自动化办公机器人")
    print("=" * 60)
    print("请选择模型")
    for i,m in enumerate(MODEL_TYPE):
        print(f"{i+1}.{m}")
    while True:
        try:
            choice = int(input("请输入有效序号,输入其他内容均无效：\n").strip())
            if choice not in [1,len(MODEL_TYPE)]:
                print("序号不存在，请重新输入")
                continue
            model_name = MODEL_TYPE[choice - 1]
            logger.info(f"已选择模型：{model_name}")
            break
        except Exception:
            logger.error("输入无效内容，请重新输入")
            continue
    # 功能菜单
    menu = """
    请选择功能：
    1.AI生成Word文档
    2.AI分析Excel数据
    3.AI发送个性化邮件
    4.批量生成Word文档
    0：退出"""
    while True:
        print(menu)
        choice = input("输入功能编号：").strip()
        if choice == "0":
            print("退出程序，再见！")
            break
        elif choice == "1":
            topic = input("输入文档主题：").strip()
            if topic:
                ai_generate_word(topic,model_name)
            else:
                print("主题不能为空")
        elif choice == "2":
            file_path = input("输入Excel文件路径：").strip()
            if file_path:
                ai_analyze_excel(model_name,file_path)
            else:
                print("文件路径不能为空")
        elif choice == "3":
            to_email = input("输入收件人邮箱：").strip()
            subject = input("输入邮件主题：").strip()
            if to_email and subject:
                ai_send_email(model_name,to_email,subject)
            else:
                print("邮箱和主题都不能为空")
        elif choice == "4":
            print("批量生成Word文档（输入q结束）")
            topics = []
            while True:
                topic = input(f"输入第{len(topics) + 1}文档主题：\n").strip()
                if topic.lower() == "q":
                    break
                if topic:
                    topics.append(topic)
            if topics:
                batch_process_documents(model_name,topics)
        else:
            print("无效选择，请重新输入!")
            input("\n按回车继续...")
            os.system("cls" if os.name == "nt" else "clear")
if __name__ == "__main__":
    main()